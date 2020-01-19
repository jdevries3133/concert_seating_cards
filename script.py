'''
The purpose of this program is to input rosters for multiple ensembles, and to output a seating chart,
plus a word document which can be printed onto index cards; in order to create a label for each student's
seat.
'''

from docx import Document

docxCB = Document('concertBand.docx')

docxWE = Document('windEnsemble.docx')

# todo Populate a 3d array as follows: [ensemble[section[individual]]]

# convert listt of paragraph pointers to listt of text
cbList = []
for y in docxCB.paragraphs:
    s = y.text
    cbList.append(s)

weList = []
for y in docxWE.paragraphs:
    s = y.text
    weList.append(s)


# make a listt of lists of sections [[section name, student, student....], [section name, student, student...]]
def createSectionArray(bandList):
    row = []
    array = []
    for index, item in enumerate(bandList):
        if item == '' and row == []:
            pass
        elif item == '':
            app = row.copy()
            array.append(app)
            del row[0:len(row)]
        else:
            row.append(item)
    app = row.copy()
    array.append(app)
    return array


# Run function above
cbArray, weArray = createSectionArray(cbList), createSectionArray(weList)


# goal = {Flutes: ['student', 'student', 'student', 'student'], Clarinets: ['student', 'student']}
def popSectionTitle(array):
    output = {}
    for index, listt in enumerate(array):
        if len(listt) >= 2:  # only accept lists with students in them
            sectionTitle = listt.pop(0)
            output.update({sectionTitle: listt})

        else:  # delete empty sections if necessary
            del array[index]

    return output


cbDictionary, weDictionary = popSectionTitle(cbArray), popSectionTitle(weArray)

# combine cbArray and weArray into goal 3d array
masterDictionary = {
    'Concert Band': cbDictionary,
    'Wind Ensemble': weDictionary
}


# todo Give the user the opportunity to alter the order of the students in each section


# prints all ensembles in the master dictionary
def displayEnsembles(masterD):
    print(''.center(100, '='))
    for key, value in masterD.items():
        print(f'Here is the {key}:\n')
        for item in value.items():
            print(item)
        print('\n')
    print(''.center(100, '='))


def selectSection(dict):
    print(f'Please enter the name of the ensemble you would like to edit:')
    usrKeyEnsemble = input().title()
    try:
        ensembleChoice = dict.get(usrKeyEnsemble)
    except AttributeError:
        pass

    print(f'Please enter the name of the section you would like to edit:')
    usrKeySection = input().title()
    try:
        sectionChoice = ensembleChoice.get(usrKeySection)
    except AttributeError:
        sectionChoice = None

    if sectionChoice == None:  # If user input returns nothing, try again
        print('Oops! An error occured. Let\'s try again!\nMake sure you type in the ensemble and section exactly'
              'as they appear above.')
        return selectSection(dict)
    else:
        return sectionChoice, usrKeySection, usrKeyEnsemble


def validate(userInput, relevantSection):
    userInputList = list(userInput)
    output = []
    # check if the user entered straight numbers with no spaces. accept only if len(section) > 11
    # (therefore max possible index = 10)
    try:
        for char in userInputList:  # error if string contains any non-integers
            int(char)
        if 11 > len(relevantSection):  # accept if there are few enough indices
            output = userInputList
            error = None
        else:  # reject if there are too many indices
            output = 'error'
            error = 'There are ten or more indicies in this section.\n' \
                    'Therefore, you must put a space between each index so that the program can differentiate\n' \
                    '"11" from "1, 1."\n'
            return output, error
    except:
        for index, char in enumerate(userInputList):
            try:  # check for two digit number
                int(char)
                int(userInputList[index + 1])
                dgt1 = char
                dgt2 = userInputList.pop(index + 1)
                output.append(dgt1 + dgt2)
            except:
                try:  # check for one digit number
                    int(char)
                    output.append(char)
                except:
                    pass

    # check for duplicates
    duplicates = []
    for x in range(len(relevantSection)):
        counter = 0
        for strIndex in output:
            index = int(strIndex)
            if index == x:
                counter = counter + 1
            else:
                pass
            if counter >= 2:
                duplicates.append(index)

    if len(duplicates) >= 1:
        return 'error', 'Error: Duplicate Indicies\nPlease try again:'



    # test if any indices are out of range
    try:
        for strIndex in output:
            test = int(strIndex)
            relevantSection.__getitem__(test)
    except IndexError:
        return 'error', 'Error: some indexes are out of range. Don\'t forget zero!\n' \
                        'Please try again:'
    except:
        pass

    if len(output) != len(relevantSection):  # check that len(output) == len(relevantSection)

        error = (f'The number of indicies you entered does not match the length of the listt of students.'
                 f'Please try again:')
        return 'error', error



    else:
        error = None
        return output, error


def editSequence():  # allows user to alter GLOBAL MASTER DICTIONARY
    global masterDictionary
    import pprint
    sectionStudents, sectionKey, ensembleKey = selectSection(masterDictionary)

    print(f'''
    Instructions:

    The number to the left of each student is that student's index. To revise the order of the listt, enter the 
    indicies in the order you would like for them to appear. For example, if you see [(1, Nick), (2, James), (3, Billy);
    and you want the section to be seated: Billy, Nick, James, simply enter \"3 1 2.\"

    Press any key to proceed.
    ''')
    input()

    # generate enumerated section listt
    tup = []
    for x in enumerate(sectionStudents):
        tup.append(x)
    insert = tuple(tup)

    print(f'Here is the current section order, seated from left to right:\n')
    pprint.pprint(insert)
    print(f'\nPlease enter the order in which you want the students to appear:')

    while True:
        order = input()
        validatedOrder, errorDescriptor = validate(order, sectionStudents)
        while validatedOrder == 'error':  # no spaces, or non-matching length
            print(errorDescriptor)
            newOrder = input()
            validatedOrder, errorDescriptor = validate(newOrder, sectionStudents)

        # reorder listt in masterDictionary
        updateList = []
        for item in validatedOrder:
            index = int(item)
            app = sectionStudents[index]
            updateList.append(app)

        sectionStudents = updateList
        masterDictionary[ensembleKey][sectionKey] = updateList
        a = masterDictionary[ensembleKey][sectionKey]


        tup = []
        for x in enumerate(a):
            tup.append(x)
        insert = tuple(tup)

        print(f'Here is the updated order:\n')
        pprint.pprint(insert)
        print('Would you like to revise this section again? (y/n)')
        yn = input()
        # user decides whether they want to make more changes to the same section        yn = input().lower()
        loop = True
        while loop == True:
            if yn == 'y':
                print('Please re-enter the order in which you want the students to appear.')
                loop = False
            elif yn == 'n':
                print('Press enter to review all ensembles again.')
                input()
                displayEnsembles(masterDictionary)
                return
            else:
                print('Please enter \"Y\" for yes, or \"N\" for no.')
                yn = input().lower()

displayEnsembles(masterDictionary)
print(f'''
{'Hello!'.center(100)}

Welcome to the seating card wizard. In this portion of the program, you will have the opportunity to adjust the
seating order of each section. Above, you can see the current order in which each section will be seated.


Would you like to make any changes (y/n)
''')

yn = input().lower()
while True:  # user decides whether they want to make changes
    if yn == 'y':
        editSequence()
        print('Would you like to make any additional changes? (y/n)')
        yn = input().lower()
    elif yn == 'n':
        break
    else:
        print('Please enter \"Y\" for yes, or \"N\" for no.')
        yn = input().lower()

# todo Process data so that [ensemble[section[individual]]] ----> row[seat[individualSSSS]]

# break sections into default sub-section:
breaks = {
    'Flute': 3,
    'Clarinet': 3,
    'Alto Sax': 2,
    'Trumpet': 3,
    'Trombone': 2,
    # Note: all other sections will be all together
}

seatingDict = {}
for ensembleName, ensembleDictionary in masterDictionary.items():
    subSectionsforSeating = {}
    for sectionName in ensembleDictionary.keys():
        breaks.setdefault(sectionName, 1)

    for section, divisor in breaks.items():

        for sectionName, sectionStudents in ensembleDictionary.items():

            if sectionName == section:

                breaks.setdefault(sectionName, 0)
                if divisor == 3:
                    split1 = len(sectionStudents) / divisor
                    split2 = split1 * 2
                    subsection1 = []
                    subsection2 = []
                    subsection3 = []
                    for index, student in enumerate(sectionStudents):
                        if split1 > index:
                            subsection1.append(student)
                        if split1 < index < split2:
                            subsection2.append(student)
                        if split2 < index:
                            subsection3.append(student)
                    app1, app2, app3 = subsection1.copy(), subsection2.copy(), subsection3.copy()
                    if len(app3) == 0:
                        subSectionsforSeating.update({sectionName: [app1, app2]})
                    else:
                        subSectionsforSeating.update({sectionName: [app1, app2, app3]})

                elif divisor == 2:
                    split = len(sectionStudents) / divisor
                    subsection1 = []
                    subsection2 = []
                    for index, student in enumerate(sectionStudents):
                        if split > index:
                            subsection1.append(student)
                        if split <= index:
                            subsection2.append(student)
                    app1, app2 = subsection1.copy(), subsection2.copy()
                    subSectionsforSeating.update({sectionName: [app1, app2]})

                elif divisor == 1:
                    app = sectionStudents.copy()
                    subSectionsforSeating.update({sectionName: app})

                update = {ensembleName: subSectionsforSeating}
                seatingDict.update(update.copy())


# Now, make the stage data structure
combinedBand = {
    'Row 1':
        seatingDict['Wind Ensemble']['Clarinet'][0]
        + seatingDict['Concert Band']['Clarinet'][0]
        + seatingDict['Wind Ensemble']['Oboe']
        + seatingDict['Concert Band']['Oboe']
        + seatingDict['Concert Band']['Flute'][0]
        + seatingDict['Wind Ensemble']['Flute'][0]
        ,
    'Row 2':
        seatingDict['Wind Ensemble']['Clarinet'][1]
        + seatingDict['Wind Ensemble']['Clarinet'][2]
        + seatingDict['Concert Band']['Clarinet'][1]
        + seatingDict['Concert Band']['Clarinet'][2]
        + seatingDict['Concert Band']['Alto Sax'][1]
        + seatingDict['Wind Ensemble']['Alto Sax'][1]
        + seatingDict['Concert Band']['Alto Sax'][0]
        + seatingDict['Wind Ensemble']['Alto Sax'][0]
        ,
    'Row 3':
        seatingDict['Wind Ensemble']['Trumpet'][0]
        + seatingDict['Wind Ensemble']['Trumpet'][1]
        + seatingDict['Wind Ensemble']['Trumpet'][2]
        + seatingDict['Concert Band']['Trumpet'][0]
        + seatingDict['Concert Band']['Trumpet'][1]
        + seatingDict['Concert Band']['Trumpet'][2]
        + seatingDict['Wind Ensemble']['Bass Clarinet']
        + seatingDict['Concert Band']['Bass Clarinet']
        + seatingDict['Wind Ensemble']['Bassoon']
        + seatingDict['Wind Ensemble']['Bari Sax']
        + seatingDict['Concert Band']['Bari Sax']
        + seatingDict['Concert Band']['Horn']
        + seatingDict['Wind Ensemble']['Horn']
        ,
    'Row 4':
        seatingDict['Wind Ensemble']['Trombone'][0]
        + seatingDict['Wind Ensemble']['Trombone'][1]
        + seatingDict['Concert Band']['Trombone'][0]
        + seatingDict['Concert Band']['Trombone'][1]
        + seatingDict['Wind Ensemble']['Euphonium']
        + seatingDict['Concert Band']['Euphonium']
        + seatingDict['Wind Ensemble']['Tuba']
        + seatingDict['Concert Band']['Horn']
        + seatingDict['Wind Ensemble']['Horn']
}

concertBandSeating = {
    'Row 1':
        seatingDict['Concert Band']['Clarinet'][0]
        + seatingDict['Concert Band']['Oboe']
        + seatingDict['Concert Band']['Flute'][0]
    ,
    'Row 2':
        seatingDict['Concert Band']['Clarinet'][1]
        + seatingDict['Concert Band']['Clarinet'][2]
        + seatingDict['Concert Band']['Alto Sax'][1]
        + seatingDict['Concert Band']['Alto Sax'][0]
        + seatingDict['Concert Band']['Bass Clarinet']
        + seatingDict['Concert Band']['Tenor Sax']
        + seatingDict['Concert Band']['Horn']
    ,
    'Row 3':
        seatingDict['Concert Band']['Trumpet'][0]
        + seatingDict['Concert Band']['Trumpet'][1]
        + seatingDict['Concert Band']['Trumpet'][2]
        + seatingDict['Concert Band']['Bari Sax']
        + seatingDict['Concert Band']['Euphonium']
        + seatingDict['Concert Band']['Trombone'][0]
        + seatingDict['Concert Band']['Trombone'][1]
}

windEnsembleSeating = {
    'Row 1':
        seatingDict['Wind Ensemble']['Clarinet'][0]
        + seatingDict['Wind Ensemble']['Oboe']
        + seatingDict['Wind Ensemble']['Flute'][0]
    ,
    'Row 2':
        seatingDict['Wind Ensemble']['Clarinet'][1]
        + seatingDict['Wind Ensemble']['Clarinet'][2]
        + seatingDict['Wind Ensemble']['Bass Clarinet']
        + seatingDict['Wind Ensemble']['Bassoon']
        + seatingDict['Wind Ensemble']['Alto Sax'][1]
        + seatingDict['Wind Ensemble']['Alto Sax'][0]
        + seatingDict['Wind Ensemble']['Tenor Sax']
    ,
    'Row 3':
        seatingDict['Wind Ensemble']['Trumpet'][0]
        + seatingDict['Wind Ensemble']['Trumpet'][1]
        + seatingDict['Wind Ensemble']['Trumpet'][2]
        + seatingDict['Wind Ensemble']['Bari Sax']
        + seatingDict['Wind Ensemble']['Euphonium']
        + seatingDict['Wind Ensemble']['Trombone'][0]
        + seatingDict['Wind Ensemble']['Trombone'][1]
        + seatingDict['Wind Ensemble']['Tuba']
        + seatingDict['Wind Ensemble']['Horn']
}

#housekeeping
import copy
combinedBand = copy.deepcopy(combinedBand)
windEnsembleSeating = copy.deepcopy(windEnsembleSeating)
concertBandSeating = copy.deepcopy(concertBandSeating)

del app, app1, app2, breaks, cbArray, cbDictionary, cbList, divisor, docxCB, docxWE, \
    ensembleDictionary, ensembleName, s, section, sectionName, sectionStudents, split, split1, split2\
    , student, subSectionsforSeating, subsection1,subsection2, subsection3, update, weArray, weDictionary, weList, \
    y, yn

# todo Create, format, fill, and LAUNCH FOR THE USER TO PRINT a word document that can be printed directly onto flash
#  cards.

print('I\'d recommend the follwing lengths for your rows:\n\n'
      f'Row 1: {len(combinedBand["Row 1"])}\n'
      f'Row 2: {len(combinedBand["Row 2"])}\n'
      f'Row 3: {len(combinedBand["Row 3"])}\n'
      f'Row 4: {len(combinedBand["Row 4"])}')

#Create output data structure

output = {'Row 1': [],
          'Row 2': [],
          'Row 3': [],
          'Row 4': []
          }


# create a listt in output dictionary
for row, studentList in combinedBand.items():
    update = []
    for x in studentList:
        seat = []
        apnd = f'Combined Band: {x}'
        seat.append(apnd)
        app = seat.copy()
        update.append(app)
    output[row] = copy.deepcopy(update)


for index, (row, listt) in enumerate(concertBandSeating.items()):
    masterList = output[row]
    mlSplit = len(masterList) // 2
    iterML1 = masterList[mlSplit:0:-1]

    # split listt into part 1, part 2
    splitPoint = len(listt) // 2
    list1 = listt[splitPoint::-1]
    list2 = listt[(splitPoint + 1):]

    # reverse part 1 (done)
    # iterate through master listt in reverse from midpoint to beginning
    for index, subList in enumerate(masterList[mlSplit::-1]):
        try:
            student = list1[index]
            app = f'Concert Band: {student}'
            subList.append(app)
        except IndexError:
            pass

    # iterate through master listt from midpoint to end
    for index, subList in enumerate(masterList[(mlSplit + 1):]):
        try:
            student = list2[index]
            app = f'Concert Band: {student}'
            subList.append(app)
        except IndexError:
            pass


for index, (row, listt) in enumerate(windEnsembleSeating.items()):
    masterList = output[row]
    mlSplit = len(masterList) // 2
    iterML1 = masterList[mlSplit:0:-1]

    # split listt into part 1, part 2
    splitPoint = len(listt) // 2
    list1 = listt[splitPoint::-1]
    list2 = listt[(splitPoint + 1):]

    # reverse part 1 (done)
    # iterate through master listt in reverse from midpoint to beginning
    for index, subList in enumerate(masterList[mlSplit::-1]):
        try:
            student = list1[index]
            app = f'Wind Ensemble: {student}'
            subList.append(app)
        except IndexError:
            pass

    # iterate through master listt from midpoint to end
    for index, subList in enumerate(masterList[(mlSplit + 1):]):
        try:
            student = list2[index]
            app = f'Wind Ensemble: {student}'
            subList.append(app)
        except IndexError:
            pass

from docx import Document
from docx.shared import Inches
from docx.shared import Pt

x = Document()
y = x.sections
outputSection = y[0]
outputSection.page_height = Inches(4)
outputSection.page_width = Inches(6)
for row, nestedList in output.items():
    for index, seat in enumerate(nestedList):
        par = x.add_paragraph()
        run = par.add_run(f'{row}, Seat {index}')
        font = run.font
        font.size = Pt(16)
        font.bold = True

        for name in seat:
            par = x.add_paragraph()
            run = par.add_run(name)
            font = run.font
            font.size = Pt(13)

        x.add_page_break()

x.save('lookKateiDidIt.docx')
